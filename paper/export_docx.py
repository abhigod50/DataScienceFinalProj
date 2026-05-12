from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


IMAGE_RE = re.compile(r"!\[(.*?)\]\((.*?)\)")
HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
ORDERED_RE = re.compile(r"^\d+\.\s+(.*)$")


def strip_inline_markup(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1", text)
    return text.strip()


def is_alignment_row(cells: list[str]) -> bool:
    if not cells:
        return False
    for cell in cells:
        token = cell.strip().replace(":", "")
        if len(token) < 3 or set(token) != {"-"}:
            return False
    return True


def parse_table_row(line: str) -> list[str]:
    stripped = line.strip().strip("|")
    return [strip_inline_markup(cell.strip()) for cell in stripped.split("|")]


def set_base_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)

    title = document.styles["Title"]
    title.font.name = "Times New Roman"
    title.font.size = Pt(18)

    for style_name, size in (("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)):
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)

    if "Caption" in document.styles:
        caption = document.styles["Caption"]
        caption.font.name = "Times New Roman"
        caption.font.size = Pt(10)


def export_markdown_to_docx(source: Path, output: Path) -> None:
    document = Document()
    set_base_styles(document)

    for section in document.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    lines = source.read_text(encoding="utf-8").splitlines()
    paragraph_buffer: list[str] = []
    table_buffer: list[str] = []
    code_buffer: list[str] = []
    in_code_block = False
    saw_title = False

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        if not paragraph_buffer:
            return
        text = strip_inline_markup(" ".join(paragraph_buffer))
        if text:
            if text.startswith("Figure "):
                caption = document.add_paragraph(text, style="Caption")
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                document.add_paragraph(text)
        paragraph_buffer = []

    def flush_code() -> None:
        nonlocal code_buffer
        if not code_buffer:
            return
        para = document.add_paragraph(style="No Spacing")
        run = para.add_run("\n".join(code_buffer))
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        code_buffer = []

    def flush_table() -> None:
        nonlocal table_buffer
        if not table_buffer:
            return

        rows = [parse_table_row(line) for line in table_buffer]
        if len(rows) >= 2 and is_alignment_row(rows[1]):
            rows.pop(1)

        if not rows:
            table_buffer = []
            return

        col_count = max(len(row) for row in rows)
        table = document.add_table(rows=1, cols=col_count)
        table.style = "Table Grid"

        for idx, value in enumerate(rows[0]):
            header_run = table.rows[0].cells[idx].paragraphs[0].add_run(value)
            header_run.bold = True

        for row_values in rows[1:]:
            row_cells = table.add_row().cells
            for idx, value in enumerate(row_values):
                row_cells[idx].text = value

        document.add_paragraph()
        table_buffer = []

    def add_image(line: str) -> bool:
        match = IMAGE_RE.fullmatch(line.strip())
        if not match:
            return False

        _, rel_path = match.groups()
        image_path = (source.parent / rel_path).resolve()
        if image_path.exists():
            document.add_picture(str(image_path), width=Inches(6.2))
        else:
            document.add_paragraph(f"[Missing image: {rel_path}]")
        return True

    for line in lines:
        stripped = line.strip()

        if in_code_block:
            if stripped.startswith("```"):
                flush_code()
                in_code_block = False
            else:
                code_buffer.append(line.rstrip())
            continue

        if table_buffer and not (stripped.startswith("|") and stripped.endswith("|")):
            flush_table()

        if stripped.startswith("```"):
            flush_paragraph()
            in_code_block = True
            code_buffer = []
            continue

        heading_match = HEADING_RE.match(stripped)
        if heading_match:
            flush_paragraph()
            level = len(heading_match.group(1))
            text = strip_inline_markup(heading_match.group(2))
            if level == 1 and not saw_title:
                document.add_paragraph(text, style="Title")
                saw_title = True
            else:
                document.add_heading(text, level=min(level, 3))
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            flush_paragraph()
            table_buffer.append(line)
            continue

        if add_image(stripped):
            flush_paragraph()
            continue

        if stripped.startswith("- "):
            flush_paragraph()
            document.add_paragraph(strip_inline_markup(stripped[2:]), style="List Bullet")
            continue

        ordered_match = ORDERED_RE.match(stripped)
        if ordered_match:
            flush_paragraph()
            document.add_paragraph(strip_inline_markup(ordered_match.group(1)), style="List Number")
            continue

        if not stripped:
            flush_paragraph()
            continue

        paragraph_buffer.append(stripped)

    flush_paragraph()
    flush_table()
    flush_code()

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def main() -> None:
    paper_dir = Path(__file__).resolve().parent
    parser = argparse.ArgumentParser(description="Export the final paper Markdown to DOCX")
    parser.add_argument("--source", type=Path, default=paper_dir / "final_paper.md")
    parser.add_argument("--output", type=Path, default=paper_dir / "final_paper.docx")
    args = parser.parse_args()

    export_markdown_to_docx(args.source.resolve(), args.output.resolve())
    print(f"Wrote DOCX to {args.output.resolve()}")


if __name__ == "__main__":
    main()